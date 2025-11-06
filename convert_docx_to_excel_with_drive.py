import os
import io
import re
import time
import requests
import mimetypes
import pathlib
import pandas as pd
from typing import Optional, Tuple, List
from google.auth.transport.requests import Request

from docx import Document

# Google API imports
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from googleapiclient.errors import HttpError
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow

# ---------------------------
# CONFIG
# ---------------------------
INPUT_DOCX = "College of Fisheries.docx"   # Change if needed
OUTPUT_XLSX = "College of Fisheries.xlsx"
DRIVE_FOLDER_NAME = "DOCX Image Uploads"
SCOPES = ["https://www.googleapis.com/auth/drive.file"]  # minimal scope for files you create

# If your document uses a different header name for the image column, add it here
PHOTO_HEADER_CANDIDATES = {"Photo", "Image", "Picture", "Photograph"}

# ---------------------------
# Google Drive helpers
# ---------------------------

def get_drive_service() -> any:
    """
    Returns an authenticated Drive API client.
    Requires credentials.json in the current directory on first run.
    """
    creds = None
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            # Attempt refresh
            try:
                creds.refresh(Request())           # RIGHT
            except Exception:
                pass
        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)
        with open("token.json", "w") as token:
            token.write(creds.to_json())
    return build("drive", "v3", credentials=creds)


def find_or_create_folder(service, name: str) -> str:
    """Find a folder by exact name or create it; return folder ID."""
    safe_name = name.replace("'", "\\'")
    q = (
        f"name = '{safe_name}' and "
        f"mimeType = 'application/vnd.google-apps.folder' and "
        f"trashed = false"
    )
    try:
        res = service.files().list(
            q=q,
            spaces="drive",
            fields="files(id, name)",
            pageSize=10,
            includeItemsFromAllDrives=False,
            supportsAllDrives=False,
            corpora="user",
        ).execute()
        files = res.get("files", [])
        if files:
            return files[0]["id"]
    except Exception:
        pass

    folder = service.files().create(
        body={"name": name, "mimeType": "application/vnd.google-apps.folder"},
        fields="id",
        supportsAllDrives=False,
    ).execute()
    return folder["id"]



def set_public_anyone_reader(service, file_id: str) -> None:
    """
    Makes a file publicly readable.
    """
    try:
        service.permissions().create(
            fileId=file_id,
            body={"type": "anyone", "role": "reader"},
        ).execute()
    except HttpError as e:
        # If permission already exists or rate-limited, try a light backoff
        if e.resp.status in (403, 429, 500, 503):
            time.sleep(2)
            service.permissions().create(
                fileId=file_id,
                body={"type": "anyone", "role": "reader"},
            ).execute()
        else:
            raise


def upload_image_bytes(service, folder_id: str, name_hint: str, data: bytes, mime: Optional[str]) -> str:
    """
    Upload raw image bytes to Drive, return the public view URL.
    """
    if not mime:
        mime = "application/octet-stream"
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime, resumable=False)
    metadata = {"name": name_hint, "parents": [folder_id]}
    file = service.files().create(body=metadata, media_body=media, fields="id").execute()
    file_id = file["id"]
    set_public_anyone_reader(service, file_id)
    # Public view URL that renders the image
    return f"https://drive.google.com/uc?export=view&id={file_id}"


def upload_image_file(service, folder_id: str, file_path: str) -> str:
    """
    Upload a local image file to Drive, return the public view URL.
    """
    mime, _ = mimetypes.guess_type(file_path)
    with open(file_path, "rb") as f:
        data = f.read()
    name_hint = pathlib.Path(file_path).name
    return upload_image_bytes(service, folder_id, name_hint, data, mime)


# ---------------------------
# DOCX table parsing helpers
# ---------------------------

def cell_text(cell) -> str:
    paras = [p.text.strip() for p in cell.paragraphs]
    joined = "\n".join([p for p in paras if p is not None])
    return joined.strip()


def first_hyperlink_url(cell) -> Optional[str]:
    """
    Return the first external hyperlink URL found in the cell, if any.
    Handles both w:hyperlink with r:id and field-code HYPERLINK cases.
    """
    try:
        tc = cell._tc
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        # direct hyperlinks
        for h in tc.xpath(".//w:hyperlink[@r:id]", namespaces=ns):
            rId = h.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rId and rId in cell.part.rels:
                rel = cell.part.rels[rId]
                if getattr(rel, "is_external", False):
                    return rel.target_ref

        # field code hyperlinks
        for instr in tc.xpath(".//w:instrText", namespaces=ns):
            t = instr.text or ""
            if "HYPERLINK" in t:
                m = re.search(r'HYPERLINK\s+"([^"]+)"', t)
                if m:
                    return m.group(1)
    except Exception:
        return None
    return None


def first_embedded_image_bytes(cell) -> Optional[Tuple[bytes, str]]:
    """
    Returns (blob, mime) of the first embedded image in a cell (supports inline & anchor images).
    """
    try:
        # Loop over all images in cell's part relationships
        rels = cell.part.rels
        for rel in rels:
            target = rels[rel].target_ref
            # Skip non-image relationships
            if not hasattr(rels[rel].target_part, "content_type"):
                continue
            content_type = rels[rel].target_part.content_type
            if content_type.startswith("image/"):
                image_part = rels[rel].target_part
                blob = image_part.blob
                mime = content_type
                return blob, mime
    except:
        pass
    return None



def download_url_bytes(url: str) -> Optional[Tuple[bytes, str]]:
    try:
        r = requests.get(url, timeout=20)
        r.raise_for_status()
        data = r.content
        mime = r.headers.get("Content-Type", None)
        return data, mime
    except Exception:
        return None


# ---------------------------
# Main conversion
# ---------------------------

def convert_docx_to_excel_with_drive(
    input_docx: str,
    output_xlsx: str,
    folder_name: str,
    photo_header_candidates: set = PHOTO_HEADER_CANDIDATES,
):
    # Load docx
    doc = Document(input_docx)
    if not doc.tables:
        raise ValueError("No tables found in the DOCX.")

    # Choose a table. Heuristic: first table with at least 3 columns, else first table
    table = None
    for t in doc.tables:
        if len(t.columns) >= 3:
            table = t
            break
    if table is None:
        table = doc.tables[0]

    # Build headers
    first_row = table.rows[0]
    headers = [cell_text(c) for c in first_row.cells]

    # Detect if first row is header-like
    header_like = sum(1 for h in headers if h.strip()) >= max(1, int(len(headers) * 0.6))
    if header_like:
        data_rows = table.rows[1:]
    else:
        # Generate generic headers
        headers = [f"Column {i+1}" for i in range(len(table.rows[0].cells))]
        data_rows = table.rows[:]

    # Identify the Photo column index if present
    photo_idx = None
    for i, h in enumerate(headers):
        if h.strip().lower() in {c.lower() for c in photo_header_candidates}:
            photo_idx = i
            break

    # Authenticate Drive and prepare folder
    drive = get_drive_service()
    folder_id = find_or_create_folder(drive, folder_name)

    rows_out: List[dict] = []
    for r_i, row in enumerate(data_rows, start=1):
        values = []
        for c_i, cell in enumerate(row.cells):
            values.append(cell_text(cell))

        # If Photo column exists, upload and replace with public link
        if photo_idx is not None and photo_idx < len(row.cells):
            photo_cell = row.cells[photo_idx]
            public_link = None

            # 1) Embedded image in the cell
            embedded = first_embedded_image_bytes(photo_cell)
            if embedded:
                blob, mime = embedded
                name_hint = f"row{r_i}_photo"
                public_link = upload_image_bytes(drive, folder_id, name_hint, blob, mime)

            # 2) Hyperlink present. If yes, download it then upload to Drive
            if not public_link:
                url = first_hyperlink_url(photo_cell)
                if url:
                    fetched = download_url_bytes(url)
                    if fetched:
                        blob, mime = fetched
                        name_hint = f"row{r_i}_photo_from_link"
                        public_link = upload_image_bytes(drive, folder_id, name_hint, blob, mime)
                    else:
                        # Keep the original URL if download failed
                        public_link = url

            # 3) If neither image nor link, leave text as-is
            if public_link:
                values[photo_idx] = public_link

        # Pack row dict with headers
        row_dict = {}
        for i, h in enumerate(headers):
            row_dict[h if h else f"Column {i+1}"] = values[i] if i < len(values) else ""
        rows_out.append(row_dict)

    # Save to Excel
    df = pd.DataFrame(rows_out, columns=headers)
    df.to_excel(output_xlsx, index=False)
    print(f"Saved Excel: {output_xlsx}")
    print(f"Images uploaded to Drive folder: {DRIVE_FOLDER_NAME}")


if __name__ == "__main__":
    convert_docx_to_excel_with_drive(INPUT_DOCX, OUTPUT_XLSX, DRIVE_FOLDER_NAME)